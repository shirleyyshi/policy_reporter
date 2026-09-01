from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.response import Response
from rest_framework import status as http_status
from django.http import HttpResponse
from django.db import IntegrityError, transaction
from .models import CentralPolicy, LocalPolicy
from io import BytesIO
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from collections import defaultdict
import datetime
from openai import OpenAI  # 使用 OpenAI SDK 兼容 DeepSeek
from django.conf import settings
from django.contrib.auth.models import User
from django.contrib.auth.password_validation import validate_password
from django.core.exceptions import ValidationError

# DeepSeek API 配置（从 settings 读取，密钥外置到 .env）
openai_client = OpenAI(api_key=settings.DEEPSEEK_API_KEY, base_url=settings.DEEPSEEK_BASE_URL)


# 辅助样式函数
def set_font(run, font_name="微软雅黑", font_size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_hyperlink(paragraph, url, text, font_name="微软雅黑", font_size=11):
    # 空 URL 必须降级为纯文本：Target="" 的外链关系会被 Word 判为非法文档直接拒开
    # （python-docx 宽松可打开，Word 严格校验整份文件）
    if not url:
        run = paragraph.add_run(text)
        set_font(run, font_name=font_name, font_size=font_size)
        return None
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rPr.append(color)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def call_deepseek_summarization(texts):
    """调 DeepSeek 生成摘要，失败时指数退避重试 3 次。
    Call DeepSeek for summarization with exponential-backoff retry (3 attempts).
    """
    # 没有政策内容时直接返回空，不调用 API
    if not texts:
        return ""

    # 摘要条数 = 政策条数，最多5条；1条政策就只生成1条摘要
    target_count = min(len(texts), 5)
    # 只取前 target_count 条政策内容送入 prompt
    selected_texts = texts[:target_count]

    prompt = (
            "假设你是一位财税专家，为公司管理层梳理每日日报热点。"
            f"请阅读以下{len(selected_texts)}条财税类政策内容，为每条政策生成1条摘要要点，共{target_count}条。"
            "语言正式、简洁，适合放在财税简报的开头部分，每条20-35字。"
            "不需要标题或者【今日财税热点摘要】之类的开头。"
            "每条摘要前请加一个圆点（• ），并换行显示，禁止使用星号(*)、序号（如1. 2. 3.）或其他Markdown格式符号。"
            f"必须且只能生成{target_count}条摘要，不要多生成也不要少生成。仅返回纯文本，不要多余解释。\n\n"
            + "\n\n".join(selected_texts)
    )

    # 指数退避重试（应对 DeepSeek 限流/网络抖动）
    last_exc = None
    for attempt in range(3):
        try:
            response = openai_client.chat.completions.create(
                model="deepseek-chat",
                messages=[{"role": "user", "content": prompt}],
                stream=False
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            last_exc = e
            if attempt < 2:
                import time
                delay = 1.0 * (2 ** attempt)
                import logging
                logging.getLogger(__name__).warning(
                    f"DeepSeek 摘要调用失败（第 {attempt+1}/3 次），{delay}s 后重试: {e}"
                )
                time.sleep(delay)
    # 重试耗尽，返回空摘要（不阻断导出流程）
    import logging
    logging.getLogger(__name__).error(f"DeepSeek 摘要重试耗尽: {last_exc}")
    return ""


def set_heading_font(paragraph, font_name="微软雅黑", font_size=14, bold=True):
    if paragraph.runs:
        for run in paragraph.runs:
            set_font(run, font_name=font_name, font_size=font_size, bold=bold)


def generate_docx(central, local, legal_text, output_stream, summary=None, report_date=None,
                  related_analysis=None):
    doc = Document()

    # 标题日期：优先用调用方传入的日期（前端 selectedDate），否则回退到当天
    if report_date:
        # 兼容 "2025-08-01" / "2025-08-01T00:00:00" 等格式，统一成 YYYY.MM.DD
        date_str_for_title = report_date[:10].replace("-", ".")
    else:
        date_str_for_title = datetime.datetime.now().strftime("%Y.%m.%d")

    # 摘要：若调用方传入预计算 summary（Agent 路径）则直接用；否则调 DeepSeek 生成（基线路径）
    if summary is None:
        policy_texts = [content for _, content, _, _ in central] + [content for _, content, _, _ in local]
        summary = call_deepseek_summarization(policy_texts)

    # 文档属性（Word「文件-信息」面板与搜索索引可见，与文件名口径一致）
    doc.core_properties.title = f"每日财税日报（{date_str_for_title}）"
    doc.core_properties.author = "Policy Reporter"

    # 主标题
    h1 = doc.add_heading(f"每日财税日报（{date_str_for_title}）", 0)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_heading_font(h1, font_size=20)

    # 今日热点资讯
    h_today = doc.add_heading("今日热点资讯", level=1)
    set_heading_font(h_today)

    # 摘要正文段落
    p_summary = doc.add_paragraph(summary)
    if not p_summary.runs:
        p_summary.add_run()
    set_font(p_summary.runs[0], font_size=11)

    # 最新发布的其他政策法规
    h_latest = doc.add_heading("最新发布的其他政策法规", level=1)
    set_heading_font(h_latest)

    # 中央政策
    h_central = doc.add_heading("中央政策", level=2)
    set_heading_font(h_central, font_size=12)

    type_group = defaultdict(list)
    for title, content, policy_type, source_url in central:
        type_group[policy_type].append((title, source_url))

    for policy_type, policies in type_group.items():
        p_type = doc.add_paragraph(style="List Bullet")
        run_type = p_type.add_run(f"【{policy_type}】")
        set_font(run_type, font_size=12)

        for title, url in policies:
            p_item = doc.add_paragraph(style="List Bullet 2")
            add_hyperlink(p_item, url, title)
            if p_item.runs:
                set_font(p_item.runs[0], font_size=11)

    # 地方法规（按业务分类 type 分组，与中央一致）
    h_local = doc.add_heading("地方法规", level=2)
    set_heading_font(h_local, font_size=12)

    type_group_local = defaultdict(list)
    for title, content, policy_type, source_url in local:
        type_group_local[policy_type or "综合"].append((title, source_url))

    for policy_type, policies in type_group_local.items():
        p_type = doc.add_paragraph(style="List Bullet")
        run_type = p_type.add_run(f"【{policy_type}】")
        set_font(run_type, font_size=12)

        for title, url in policies:
            p_item = doc.add_paragraph(style="List Bullet 2")
            add_hyperlink(p_item, url, title)
            if p_item.runs:
                set_font(p_item.runs[0], font_size=11)

    # 关联政策分析（Agent 路径可选：基于历史政策的延续/修订/配套分析，独立于当日摘要）
    if related_analysis:
        h_related = doc.add_heading("关联政策分析", level=1)
        set_heading_font(h_related)
        for para in related_analysis.splitlines():
            line = para.strip()
            if not line:
                continue
            p = doc.add_paragraph(line)
            run = p.runs[0] if p.runs else p.add_run()
            set_font(run, font_size=11)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 法律法规及合规资讯
    h_legal = doc.add_heading("法律法规及合规资讯", level=1)
    set_heading_font(h_legal)

    if legal_text:
        for para in legal_text.splitlines():
            p = doc.add_paragraph(para.strip())
            run = p.runs[0] if p.runs else p.add_run()
            set_font(run, font_name="微软雅黑", font_size=11)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p = doc.add_paragraph("无相关内容。")
        run = p.runs[0]
        set_font(run, font_name="微软雅黑", font_size=11)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 空一行
    doc.add_paragraph()

    # 分割线
    p_split = doc.add_paragraph("— — — — — — — — — —")
    p_split.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_split.runs:
        set_font(p_split.runs[0], font_size=10)
        p_split.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    # 个人项目署名（不涉及版权声明，仅标注来源）
    p_sign = doc.add_paragraph()
    run_sign = p_sign.add_run(
        f"本日报由 Policy Reporter 个人学习项目自动生成 · {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')} 生成"
    )
    set_font(run_sign, font_size=9)
    run_sign.font.color.rgb = RGBColor(150, 150, 150)
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_disclaim = doc.add_paragraph()
    run_disclaim = p_disclaim.add_run("内容仅供个人学习与研究参考，请以官方发布原文为准。")
    set_font(run_disclaim, font_size=9)
    run_disclaim.font.color.rgb = RGBColor(150, 150, 150)
    p_disclaim.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_stream)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def policy_list(request):
    date_str = request.query_params.get('date')
    if date_str:
        try:
            datetime.date.fromisoformat(date_str)
        except (TypeError, ValueError):
            return Response({'detail': 'date 必须是 YYYY-MM-DD 格式'}, status=http_status.HTTP_400_BAD_REQUEST)

    central_qs = CentralPolicy.objects.all()
    local_qs = LocalPolicy.objects.all()
    if date_str:
        central_qs = central_qs.filter(publish_time__date=date_str)
        local_qs = local_qs.filter(publish_time__date=date_str)
    central = list(central_qs.values('id', 'title', 'type', 'publish_time', 'source_url', 'crawled_at'))
    local = list(local_qs.values('id', 'title', 'province', 'type', 'publish_time', 'source_url', 'crawled_at'))
    for item in central:
        item['source'] = "central"
    for item in local:
        item['source'] = "local"
    return Response({'central': central, 'local': local})


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def policy_detail(request):
    """单条政策详情（含 content 全文，供详情页展示）。

    source: central/local；id: 对应表主键。
    """
    source = request.query_params.get('source')
    policy_id = request.query_params.get('id')
    if source not in ('central', 'local'):
        return Response({'detail': 'source 必须为 central 或 local'}, status=http_status.HTTP_400_BAD_REQUEST)
    if not policy_id:
        return Response({'detail': 'id 为必填项'}, status=http_status.HTTP_400_BAD_REQUEST)
    try:
        policy_id = int(policy_id)
        if policy_id <= 0:
            raise ValueError
    except (TypeError, ValueError):
        return Response({'detail': 'id 必须是正整数'}, status=http_status.HTTP_400_BAD_REQUEST)
    try:
        qs = CentralPolicy if source == 'central' else LocalPolicy
        fields = ['id', 'title', 'content', 'type', 'publish_time', 'source_url', 'crawled_at']
        if source == 'local':
            fields.insert(3, 'province')
        policy = qs.objects.filter(id=policy_id).values(*fields).first()
    except Exception:
        return Response({'detail': '政策查询失败'}, status=http_status.HTTP_500_INTERNAL_SERVER_ERROR)
    if policy is None:
        return Response({'detail': '政策不存在'}, status=http_status.HTTP_404_NOT_FOUND)
    policy['source'] = source
    return Response(policy)



@api_view(['POST'])
@permission_classes([IsAuthenticated])
def policy_export(request):
    selected = request.data.get('selected_ids', [])
    legal_text = request.data.get('legal_text', '')
    report_date = request.data.get('date')  # 前端 selectedDate，用于标题日期

    if not isinstance(selected, list):
        return Response({'detail': 'selected_ids 必须是数组'}, status=http_status.HTTP_400_BAD_REQUEST)
    if not isinstance(legal_text, str):
        return Response({'detail': 'legal_text 必须是字符串'}, status=http_status.HTTP_400_BAD_REQUEST)
    if report_date:
        try:
            datetime.date.fromisoformat(str(report_date))
        except (TypeError, ValueError):
            return Response({'detail': 'date 必须是 YYYY-MM-DD 格式'}, status=http_status.HTTP_400_BAD_REQUEST)

    central_ids = []
    local_ids = []
    for index, item in enumerate(selected):
        if not isinstance(item, dict):
            return Response({'detail': f'selected_ids[{index}] 必须是对象'}, status=http_status.HTTP_400_BAD_REQUEST)
        source = item.get('source')
        policy_id = item.get('id')
        if source not in ('central', 'local'):
            return Response({'detail': f'selected_ids[{index}].source 无效'}, status=http_status.HTTP_400_BAD_REQUEST)
        if isinstance(policy_id, bool):
            return Response({'detail': f'selected_ids[{index}].id 必须是正整数'}, status=http_status.HTTP_400_BAD_REQUEST)
        try:
            policy_id = int(policy_id)
        except (TypeError, ValueError):
            return Response({'detail': f'selected_ids[{index}].id 必须是正整数'}, status=http_status.HTTP_400_BAD_REQUEST)
        if policy_id <= 0:
            return Response({'detail': f'selected_ids[{index}].id 必须是正整数'}, status=http_status.HTTP_400_BAD_REQUEST)
        (central_ids if source == 'central' else local_ids).append(policy_id)

    legal_text = legal_text.strip()

    central = CentralPolicy.objects.filter(id__in=central_ids).values_list('title', 'content', 'type', 'source_url')
    local = LocalPolicy.objects.filter(id__in=local_ids).values_list('title', 'content', 'type', 'source_url')

    doc_io = BytesIO()
    generate_docx(central, local, legal_text, doc_io, report_date=report_date)
    doc_io.seek(0)
    return HttpResponse(doc_io.read(),
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        headers={'Content-Disposition': 'attachment; filename="政策日报.docx"'})


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def policy_counts(request):
    central_count = CentralPolicy.objects.count()
    local_count = LocalPolicy.objects.count()
    return Response({
        "central_count": central_count,
        "local_count": local_count
    })


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def policy_dates(request):
    """返回每个发布日期的政策条数（中央+地方合并），供日期选择器着色。
    示例：{"2026-08-25": 4, "2026-08-24": 7}
    """
    from django.db.models import Count
    from django.db.models.functions import TruncDate

    counts = {}
    for model in (CentralPolicy, LocalPolicy):
        rows = (model.objects
                .annotate(day=TruncDate('publish_time'))
                .values('day')
                .annotate(n=Count('id')))
        for row in rows:
            if row['day']:
                key = row['day'].isoformat()
                counts[key] = counts.get(key, 0) + row['n']
    return Response(counts)


@api_view(['POST'])
@permission_classes([AllowAny])
def register(request):
    """创建普通用户账号；注册接口不授予管理员权限。"""
    username = str(request.data.get('username', '')).strip()
    password = str(request.data.get('password', ''))
    if not username or not password:
        return Response({'detail': '用户名和密码不能为空'}, status=http_status.HTTP_400_BAD_REQUEST)
    if len(username) > 150:
        return Response({'detail': '用户名不能超过 150 个字符'}, status=http_status.HTTP_400_BAD_REQUEST)
    if User.objects.filter(username=username).exists():
        return Response({'detail': '用户名已存在'}, status=http_status.HTTP_400_BAD_REQUEST)
    try:
        validate_password(password)
    except ValidationError as exc:
        return Response({'detail': '；'.join(exc.messages)}, status=http_status.HTTP_400_BAD_REQUEST)
    try:
        with transaction.atomic():
            User.objects.create_user(username=username, password=password)
    except IntegrityError:
        return Response({'detail': '用户名已存在'}, status=http_status.HTTP_409_CONFLICT)
    return Response({'detail': '注册成功'}, status=http_status.HTTP_201_CREATED)


@api_view(['GET'])
@permission_classes([AllowAny])
def health(request):
    """健康检查（匿名可访问，供负载均衡/容器探活）。

    只暴露存活状态与 DB 连通性，不含版本、数量等细节。
    """
    from django.db import connection
    try:
        with connection.cursor() as cursor:
            cursor.execute('SELECT 1')
        db_ok = True
    except Exception:
        db_ok = False
    return Response(
        {'status': 'ok' if db_ok else 'degraded', 'db': db_ok},
        status=http_status.HTTP_200_OK if db_ok else http_status.HTTP_503_SERVICE_UNAVAILABLE,
    )


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def me(request):
    """返回当前登录用户信息"""
    return Response({
        'id': request.user.id,
        'username': request.user.username,
        'email': request.user.email,
    })
